import re
import uuid
import subprocess
import pathlib
from docx import Document
import os
import json
from docx.oxml.ns import nsmap, qn
from docx.text.paragraph import Paragraph
import win32com.client as win32

def convert_doc_to_docx(folder):
    """
    doc -> docx 批量转换
    使用 win32com.client 调用 Word 应用进行转换。
    """
    word = win32.Dispatch("Word.Application")
    word.Visible = False  # 不显示 Word 窗口
    for filename in os.listdir(folder):
        if filename.lower().endswith(".doc") and not filename.lower().endswith(".docx"):
            doc_path = os.path.join(folder, filename)
            # 修正这一行
            docx_path = os.path.splitext(doc_path)[0] + ".docx"
            print(f"正在转换: {doc_path} -> {docx_path}")
            try:
                # 注意 Open 是大写的
                doc = word.Documents.Open(doc_path)
                doc.SaveAs(docx_path, FileFormat=16)  # 16 是 wdFormatXMLDocument
                doc.Close()
                os.remove(doc_path)  # 删除原始 .doc 文件
                print(f"已删除原始文件: {doc_path}")
            except Exception as e:
                print(f"转换失败: {doc_path}, 错误: {e}")

    word.Quit()  # 关闭 Word 应用

def iter_textbox_paragraphs(doc):
    """
    生成器：遍历整个文档，找出 w:txbxContent 里的 w:p，
    yield 回 python-docx 的 Paragraph 对象。
    """
    # 1) 找到所有 <w:txbxContent> 节点
    for txbx in doc._element.xpath('.//w:txbxContent', namespaces=nsmap):
        # 2) 其中的每个 <w:p> 当作独立段落返回
        for p in txbx.xpath('.//w:p', namespaces=nsmap):
            yield Paragraph(p, doc)

def save_run_image(image_part, qid):
    """保存图片并返回文件名"""
    fmt = image_part.content_type.split("/")[-1]  # png / jpeg / x-wmf
    
    # 使用图片内容的哈希值作为唯一标识，避免重复保存
    import hashlib
    image_hash = hashlib.md5(image_part.blob).hexdigest()[:8]
    fname = f"{qid}_{image_hash}.{fmt}"
    
    out_dir = pathlib.Path(r"E:\NLP_Model\ai_edu\data\processed_data\images")
    out_dir.mkdir(exist_ok=True)
    path = out_dir / fname
    
    # 如果文件已存在，直接返回文件名，不重复保存
    if path.exists():
        print(f"图片已存在，跳过保存: {fname}")
        return fname
    
    path.write_bytes(image_part.blob)
    
    # 如果是 WMF，再转 PNG 备份
    if fmt == "x-wmf":
        png = out_dir / f"{qid}_{image_hash}.png"
        try:
            subprocess.run(
                ["magick", "convert", str(path), str(png)],
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                check=True,
            )
            print(f"已将 WMF 转换为 PNG: {png}")
            os.remove(path)
            print(f"已删除原始 WMF 文件: {path}")
            return png.name
        except subprocess.CalledProcessError as e:
            print(f"WMF 转换失败: {e}")
            return fname
    
    print("-------------------已运行save_run_image方法---------------------")
    return fname

# 在文件顶部添加全局变量
global_processed_rids = set()

def extract_images_from_runs(paragraph, doc, qid):
    global global_processed_rids
    segment = ""
    
    for run in paragraph.runs:
        xml = run._element.xml
        text = run.text
        if text:
            segment += text
        
        # 查找所有图片引用
        rids = re.findall(r'(?:r:id|r:embed)="(rId\d+)"', xml)
        for rid in rids:
            # 使用全局集合检查是否已处理过此rId
            if rid in global_processed_rids:
                # 即使已处理，也要添加占位符
                try:
                    if rid in doc.part.related_parts:
                        part = doc.part.related_parts[rid]
                        if part.content_type.startswith("image/"):
                            # 生成相同的文件名但不重复保存
                            import hashlib
                            image_hash = hashlib.md5(part.blob).hexdigest()[:8]
                            fmt = part.content_type.split("/")[-1]
                            fname = f"{qid}_{image_hash}.png" if fmt == "x-wmf" else f"{qid}_{image_hash}.{fmt}"
                            segment += f"[IMG:{fname}]"
                except:
                    pass
                continue
                
            global_processed_rids.add(rid)
            
            try:
                if rid in doc.part.related_parts:
                    part = doc.part.related_parts[rid]
                    if part.content_type.startswith("image/"):
                        fname = save_run_image(part, qid)
                        segment += f"[IMG:{fname}]"
                        print(f"处理图片: {rid} -> {fname}")
            except Exception as e:
                print(f"处理rId {rid}时出错: {e}")
    
    return segment

def parse_filename_metadata(filename):
    """从文件名中提取年份、省份、城市、学科、考试类型"""
    # 默认值，如果无法解析
    metadata = {
        "year": "0000",
        "province": "未知省份",
        "city": "未知城市",
        "subject": "未知学科",
        "exam_type": "未知考试类型"
    }
    
    # 尝试从文件名解析信息
    # 解析年份
    year_match = re.search(r'(\d{4})年', filename)
    if year_match:
        metadata["year"] = year_match.group(1)
    
    # 省份识别
    provinces = ['湖南省', '湖南', '湖北省', '湖北', '江苏省', '江苏', '浙江省', '浙江', '广东省', '广东']
    for province in provinces:
        if province in filename:
            # 标准化省份名称（确保包含"省"字）
            if '省' not in province:
                metadata["province"] = f"{province}省"
            else:
                metadata["province"] = province
            break
    
    # 城市识别
    cities = ['湘西州', '湘西', '扬州市', '扬州', '长沙市', '长沙', '武汉市', '武汉', '金华市', '金华', '南京市', '南京', '杭州市', '杭州', '广州市', '广州', '徐州市', '徐州', '苏州市', '苏州', '无锡市', '无锡', '泰州市', '泰州']
    for city in cities:
        if city in filename:
            # 标准化城市名称（确保包含"市"或"州"）
            if '市' not in city and '州' not in city:
                if '扬州' in city:
                    metadata["city"] = f"{city}州"
                else:
                    metadata["city"] = f"{city}市"
            else:
                metadata["city"] = city
            break
    
    # 学科识别
    subjects = ['数学', '语文', '英语', '物理', '化学']
    for subject in subjects:
        if subject in filename:
            metadata["subject"] = subject
            break
    
    # 考试类型识别
    exam_types = ['中考', '高考', '模拟考试', '模拟']
    for exam_type in exam_types:
        if exam_type in filename:
            if exam_type == '模拟':
                metadata["exam_type"] = '模拟考试'
            else:
                metadata["exam_type"] = exam_type
            break
    
    return metadata

def create_question_id(metadata, question_number, sub_question_number=None):
    """创建问题ID，使用完整名称而非缩写"""
    base_id = f"{metadata['year']}-{metadata['province']}-{metadata['city']}-{metadata['subject']}-{metadata['exam_type']}-{question_number:02d}"
    
    if sub_question_number is not None:
        return f"{base_id}-{sub_question_number:02d}" # 子问题题号补零
    
    return base_id

def extract_doc_content(doc_path):
    """提取文档中的文本、图片及结构化题目，并在content中按原位置插入图片占位符"""
    global global_processed_rids
    global_processed_rids = set()
    filename = os.path.basename(doc_path)
    metadata = parse_filename_metadata(filename)
    print(f"解析文件元数据: {metadata}")
    
    doc = Document(doc_path)
    questions = []
    current = None
    collecting = False
    in_answer_section = False
    current_answer_sub_number = None  # 新增：当前答案对应的小题号

    # 获取文档中所有元素的顺序（段落和表格）
    def get_document_elements():
        """获取文档中段落和表格的顺序"""
        elements = []
        for element in doc.element.body:
            if element.tag.endswith('}p'):  # 段落
                # 找到对应的段落对象
                for para in doc.paragraphs:
                    if para._element == element:
                        elements.append(('paragraph', para))
                        break
            elif element.tag.endswith('}tbl'):  # 表格
                # 找到对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        elements.append(('table', table))
                        break
        return elements

    for element_type, element in get_document_elements():
        if element_type == 'paragraph':
            para = element
            raw = para.text.strip()
            splits = re.split(r'(?=[（(][一二三四五六七八九十1234567890]+[）)])', raw)
            
            for seg in splits:
                seg = seg.strip()
                if seg:
                    # 检查本段是否新题号
                    m_q = re.match(r'^(\d+)\.', seg)
                    # 检查是否是子题目
                    m_sub_q = re.match(r'[\(（](\d+)[\)）]', seg)

                if m_q:
                    qid = m_q.group(1)
                    question_id = create_question_id(metadata, int(qid))
                    
                    # 修改：检查是否已存在相同题号的题目
                    existing_id = None
                    for q in questions:
                        if q["number"] == int(qid):
                            existing_id = q
                            break
                    
                    if existing_id:
                        current = existing_id
                    
                    else:
                        # 新题目开始，重置状态
                        in_answer_section = False
                        current_answer_sub_number = None
                        
                        # 关闭上一个题
                        if current:
                            questions.append(current)
                        qid = m_q.group(1)                     # 对应正则中 第一个括号 (\d+) 捕获的内容
                        question_id = create_question_id(metadata, int(qid))
                        
                        # 新题初始化
                        current = {
                            "id": question_id,
                            "number": int(qid),
                            "content": "",
                            "options": [],
                            "answers": ""  # 添加答案字段
                        }
                        questions.append(current)

                    collecting = True
                
                # 处理子题目的情况
                elif m_sub_q and current and collecting:
                    sub_qid = m_sub_q.group(1)
                    question_id = create_question_id(metadata, current["number"], int(sub_qid))
                    
                    # 先保存当前主题目
                    if current and "sub_questions" not in current:
                        # 保存主题目的内容
                        main_content = current["content"]
                        main_options = current["options"]
                        
                        # 重置主题目，添加子题目容器
                        current["content"] = main_content  # 保留主题目的题干
                        current["options"] = []
                        current["sub_questions"] = []
                    
                    existing_sub = None
                    if "sub_questions" in current:
                        for sub_q in current["sub_questions"]:
                            if sub_q["number"] == int(sub_qid):
                                existing_sub = sub_q
                                break

                    if not existing_sub:
                        # 添加新的子题目
                        sub_question = {
                            "id": question_id,
                            "number": int(sub_qid),
                            "content": "",
                            "options": [],
                            "answers": ""  # 为子题目也添加答案字段
                        }
                        current["sub_questions"].append(sub_question)
                
                # 检查是否进入答案解析部分
                if current and collecting:
                    # 识别答案解析的关键词和格式
                    answer_patterns = [
                        r'^【答案】',
                        r'^【解析】',
                        r'^【解答】',
                        r'^答案[:：]',
                        r'^解析[:：]',
                        r'^解答[:：]'
                    ]
                    
                    # 检查是否包含小题号的答案格式
                    sub_answer_patterns = [
                        r'^\(\s*(\d+)\s*\).*?【答案】',
                        r'^\(\s*(\d+)\s*\).*?【解析】',
                        r'^\(\s*(\d+)\s*\).*?【解答】',
                        r'^\(\s*(\d+)\s*\).*?答案',
                        r'^\(\s*(\d+)\s*\).*?解析',
                        r'^(\d+)\..*?【答案】',
                        r'^(\d+)\..*?【解析】',
                        r'^(\d+)\..*?【解答】'
                    ]
                    
                    # 检查是否是带小题号的答案
                    sub_answer_match = None
                    for pattern in sub_answer_patterns:
                        sub_answer_match = re.search(pattern, seg)
                        if sub_answer_match:
                            current_answer_sub_number = int(sub_answer_match.group(1)) # 提取小题号并转为整数
                            in_answer_section = True
                            break
                    
                    # 检查是否是普通的答案开始
                    if not sub_answer_match:
                        is_answer_line = any(re.search(pattern, seg) for pattern in answer_patterns) # 任意一个返回非空值，设为True
                        if is_answer_line:
                            in_answer_section = True
                            current_answer_sub_number = None  # 重置小题号
                        elif not in_answer_section and len(seg) < 10 and ('答案' in seg or '解析' in seg):
                            # 单独的"答案"或"解析"标题行
                            in_answer_section = True
                            current_answer_sub_number = None
                            continue  # 跳过标题行本身
                
                if not collecting or current is None:
                    continue

                # 按 run 遍历，构造本段 content 片段
                segment = extract_images_from_runs(para, doc, current["id"])

                # 去除纯空白段
                if not segment.strip():
                    continue

                # 判断是否为选项行
                m_opt = re.match(r'^[A-D]\.', segment.strip())

                # 确定应该添加内容到主题目还是子题目
                target = current

                if in_answer_section:
                    # 在答案解析部分，统一存储到大题的answers中
                    if current["answers"]:
                        current["answers"] += "\n" + segment
                    else:
                        current["answers"] = segment
                        
                elif m_opt:
                    # 选项内容，添加到当前活跃的题目
                    if "sub_questions" in current and current["sub_questions"]:
                        target = current["sub_questions"][-1]
                    target["options"].append({
                        "text": segment.strip()
                    })
                else:
                    # 普通题干内容 - 关键修改：根据内容判断归属
                    # 检查内容是否包含小题标记
                    sub_content_match = re.search(r'[（(](\d+)[）)]', segment)
                    
                    if sub_content_match:
                        sub_num = int(sub_content_match.group(1))
                        # 查找对应编号的子题目
                        target_found = False
                        if "sub_questions" in current:
                            for sub_q in current["sub_questions"]:
                                if sub_q["number"] == sub_num:
                                    target = sub_q
                                    target_found = True
                                    break
                        
                        # 如果没找到对应的子题目，添加到主题目
                        if not target_found:
                            target = current
                    else:
                        # 没有小题标记的内容
                        # 如果主题目内容为空，添加到主题目；否则添加到最后一个子题目
                        if not current["content"] and "sub_questions" not in current:
                            target = current
                        elif "sub_questions" in current and current["sub_questions"]:
                            target = current["sub_questions"][-1]
                        else:
                            target = current
                    
                    # 过滤内容：确保主题目不包含小题内容，小题不包含其他小题内容
                    if target == current:
                        # 主题目：移除所有小题标记的内容
                        clean_segment = re.sub(r'[（(]\d+[）)][^（(]*', '', segment).strip()
                        if clean_segment:
                            if target["content"]:
                                target["content"] += "\n" + clean_segment
                            else:
                                target["content"] = clean_segment
                    else:
                        # 子题目：只保留对应编号的内容
                        target_num = target["number"]
                        # 提取该小题对应的内容
                        pattern = rf'[（(]{target_num}[）)][^（(]*?(?=[（(]\d+[）)]|$)'
                        matches = re.findall(pattern, segment)
                        if matches:
                            clean_content = matches[0].strip()
                        else:
                            # 如果没有匹配到特定格式，但确定是该小题的内容，则使用原内容
                            clean_content = segment.strip()
                        
                        if clean_content:
                            if target["content"]:
                                target["content"] += "\n" + clean_content
                            else:
                                target["content"] = clean_content
            
        elif element_type == 'table':
            table = element
            if collecting and current is not None:
                # 提取表格内容并格式化
                table_content = "\n[表格开始]\n"
                for row_idx, row in enumerate(table.rows):
                    row_cells = []
                    for cell in row.cells:
                        cell_text = ""
                        # 处理单元格中的段落和图片
                        for cell_para in cell.paragraphs:
                            cell_segment = extract_images_from_runs(cell_para, doc, current["id"])
                            if cell_segment.strip():
                                cell_text += cell_segment.strip() + " "
                        row_cells.append(cell_text.strip())
                    table_content += " | ".join(row_cells) + "\n"
                table_content += "[表格结束]\n"
                
                # 确定应该添加内容到主题目还是子题目
                target = current
                
                if in_answer_section:
                    # 在答案解析部分，统一存储到大题的answers中
                    if current["answers"]:
                        current["answers"] += table_content
                    else:
                        current["answers"] = table_content
                else:
                    # 非答案部分，添加到content
                    if "sub_questions" in current and current["sub_questions"]:
                        target = current["sub_questions"][-1]
                    
                    if target["content"]:
                        target["content"] += table_content
                    else:
                        target["content"] = table_content

    # # 最后一题
    # if current:
    #     questions.append(current)

    unique_questions = []
    seen_numbers = set()
    for q in questions:
        if q["number"] not in seen_numbers:
            unique_questions.append(q)
            seen_numbers.add(q["number"])
    return unique_questions, metadata

def prepare_final_questions(questions):
    """准备最终的问题列表，将大题answers分配到对应小题中"""
    final_questions = []
    
    for q in questions:
        if "sub_questions" in q and q["sub_questions"]:
            # 有子问题，需要从大题answers分配到各小题
            all_answers_content = q["answers"] if q["answers"] else ""
            
            # 初始化子题目
            processed_sub_questions = []
            for sub_q in q["sub_questions"]:
                processed_sub_questions.append({
                    "id": sub_q["id"],
                    "number": sub_q["number"],
                    "content": sub_q["content"],
                    "options": sub_q["options"].copy() if sub_q["options"] else [],
                    "answers": ""
                })
            
            # 从大题answers中提取各小题的答案
            if all_answers_content:
                # 1. 优先处理【小问X详解】格式
                for sub_q in processed_sub_questions:
                    sub_number = sub_q["number"]
                    detail_pattern = rf'【小问{sub_number}详解】(.*?)(?=【小问\d+详解】|【点睛】|$)'
                    detail_match = re.search(detail_pattern, all_answers_content, re.DOTALL)
                    if detail_match:
                        sub_q["answers"] = detail_match.group(1).strip()
                        continue
                
                # 2. 如果没有【小问X详解】格式，则综合提取各部分内容
                if not any(sub_q["answers"] for sub_q in processed_sub_questions):
                    for sub_q in processed_sub_questions:
                        sub_number = sub_q["number"]
                        answer_parts = []
                        
                        # 提取【答案】部分对应的小题内容
                        answer_section = re.search(r'【答案】(.*?)(?=【解析】|【分析】|【详解】|【点睛】|$)', all_answers_content, re.DOTALL)
                        if answer_section:
                            answer_content = answer_section.group(1)
                            pattern = rf'[（(]{sub_number}[）)](.*?)(?=[（(]\d+[）)]|$)'
                            match = re.search(pattern, answer_content, re.DOTALL)
                            if match:
                                sub_answer = match.group(1).strip().rstrip('；;')
                                # 清理可能的重复标记
                                sub_answer = re.sub(r'【答案】.*?$', '', sub_answer, flags=re.DOTALL).strip()
                                if sub_answer:
                                    answer_parts.append(f"【答案】（{sub_number}）{sub_answer}")
                        
                        # 提取【分析】部分对应的小题内容
                        analysis_section = re.search(r'【分析】(.*?)(?=【详解】|【点睛】|$)', all_answers_content, re.DOTALL)
                        if analysis_section:
                            analysis_content = analysis_section.group(1)
                            pattern = rf'[（(]{sub_number}[）)](.*?)(?=[（(]\d+[）)]|$)'
                            match = re.search(pattern, analysis_content, re.DOTALL)
                            if match:
                                sub_analysis = match.group(1).strip().rstrip('；;')
                                # 清理可能的重复标记
                                sub_analysis = re.sub(r'【分析】.*?$', '', sub_analysis, flags=re.DOTALL).strip()
                                if sub_analysis:
                                    answer_parts.append(f"【分析】（{sub_number}）{sub_analysis}")
                        
                        # 提取【详解】部分对应的小题内容
                        detail_section = re.search(r'【详解】(.*?)(?=【点睛】|$)', all_answers_content, re.DOTALL)
                        if detail_section:
                            detail_content = detail_section.group(1)
                            pattern = rf'[（(]{sub_number}[）)](.*?)(?=[（(]\d+[）)]|$)'
                            match = re.search(pattern, detail_content, re.DOTALL)
                            if match:
                                sub_detail = match.group(1).strip().rstrip('；;')
                                # 清理可能的重复标记和多余内容
                                sub_detail = re.sub(r'【详解】.*?$', '', sub_detail, flags=re.DOTALL).strip()
                                if sub_detail:
                                    answer_parts.append(f"【详解】（{sub_number}）{sub_detail}")
                        
                        # 组合所有部分
                        if answer_parts:
                            sub_q["answers"] = "\n".join(answer_parts)
            
            # 提取【点睛】部分作为大题答案
            main_answer = ""
            point_match = re.search(r'【点睛】(.*?)$', all_answers_content, re.DOTALL)
            if point_match:
                main_answer = point_match.group(0).strip()
            
            # 组装最终题目
            sorted_sub_questions = sorted(processed_sub_questions, key=lambda x: x["number"])
            
            final_questions.append({
                "id": q["id"],
                "number": q["number"],
                "content": q["content"],
                "options": q["options"],
                "sub_questions": sorted_sub_questions,
                "answers": main_answer
            })
        else:
            # 没有子问题，大题answers不分配，直接保留
            final_questions.append({
                "id": q["id"],
                "number": q["number"],
                "content": q["content"],
                "options": q["options"],
                "answers": q["answers"]
            })
    
    return final_questions


def process_document(doc_path, output_dir):
    """处理单个文档并保存结果"""
    print(f"正在处理文档: {doc_path}")
    questions, metadata = extract_doc_content(doc_path)
    
    # 准备最终的问题列表（扁平化子问题）
    final_questions = prepare_final_questions(questions)
    
    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    
    # 保存JSON文件，按照指定格式命名
    output_filename = f"{metadata['year']}-{metadata['province']}-{metadata['city']}-{metadata['subject']}-{metadata['exam_type']}.json"
    output_path = os.path.join(output_dir, output_filename)
    
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(final_questions, f, indent=2, ensure_ascii=False)
    
    print(f"已保存文件: {output_path}")
    return output_path

def process_documents_in_folder(folder_path, output_dir):
    """处理文件夹中所有docx文件"""
    processed_files = []
    
    for filename in os.listdir(folder_path):
        if filename.lower().endswith(".docx"):
            doc_path = os.path.join(folder_path, filename)
            output_path = process_document(doc_path, output_dir)
            processed_files.append((filename, output_path))
    
    return processed_files

if __name__ == "__main__":
    # 设置输入和输出目录
    input_dir = r"E:\NLP_Model\ai_edu\data\math_answer"
    output_dir = r"E:\NLP_Model\ai_edu\data\processed_data\answers"
    
    # 测试特定文件
    test_file = "2023年浙江省杭州市中考数学真题（解析版）.docx"
    test_file_path = os.path.join(input_dir, test_file)
    
    if os.path.exists(test_file_path):
        print(f"开始测试文件: {test_file}")
        print("=" * 60)
        
        # 处理单个文件并显示详细信息
        try:
            questions, metadata = extract_doc_content(test_file_path)
            
            print(f"\n文件元数据解析结果:")
            for key, value in metadata.items():
                print(f"  {key}: {value}")
            
            print(f"\n提取的题目数量: {len(questions)}")
            
            # 显示每个题目的基本信息
            for i, q in enumerate(questions, 1):
                print(f"\n题目 {i}:")
                print(f"  ID: {q['id']}")
                print(f"  题号: {q['number']}")
                print(f"  内容长度: {len(q['content'])} 字符")
                print(f"  选项数量: {len(q['options'])}")
                
                # 显示子题目信息
                if "sub_questions" in q:
                    print(f"  子题目数量: {len(q['sub_questions'])}")
                    for j, sub_q in enumerate(q['sub_questions']):
                        print(f"    子题 {j+1}: ID={sub_q['id']}, 内容长度={len(sub_q['content'])}, 选项数={len(sub_q['options'])}")
                
                # 显示内容预览（前100字符）
                content_preview = q['content'][:100].replace('\n', ' ')
                print(f"  内容预览: {content_preview}...")
            
            # 准备最终输出
            final_questions = prepare_final_questions(questions)
            print(f"\n扁平化后的题目数量: {len(final_questions)}")
            
            # 保存文件
            output_path = process_document(test_file_path, output_dir)
            print(f"\n处理完成，输出文件: {output_path}")
            
            # 验证保存的文件
            with open(output_path, 'r', encoding='utf-8') as f:
                saved_data = json.load(f)
            print(f"验证: 保存的JSON文件包含 {len(saved_data)} 个题目")
            
        except Exception as e:
            print(f"处理文件时出错: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"测试文件不存在: {test_file_path}")
        print("\n可用的文件:")
        for filename in os.listdir(input_dir):
            if filename.lower().endswith(".docx"):
                print(f"  {filename}")
    
    print("\n" + "=" * 60)
    print("测试完成")